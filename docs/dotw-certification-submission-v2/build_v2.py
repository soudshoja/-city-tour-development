"""
Build XML CERTIFICATION TEST PLAN_FILLED_v2.docx
April 2026 re-run results: 19 PASS / 2 SKIP / 0 FAIL
"""

import copy
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
import lxml.etree as etree

SRC = "C:/Users/User/OneDrive - City Travelers/soud-laravel/docs/dotw-certification-submission/XML CERTIFICATION TEST PLAN_FILLED.docx"
BLANK = "C:/Users/User/OneDrive - City Travelers/soud-laravel/docs/Dotwcert/XML CERTIFICATION TEST PLAN_.docx"
DEST = "C:/Users/User/OneDrive - City Travelers/soud-laravel/docs/dotw-certification-submission-v2/XML CERTIFICATION TEST PLAN_FILLED_v2.docx"

# ── April 3 2026 results ─────────────────────────────────────────────────────
# Keyed by table row index in the filled doc (stale) mapping.
# stale filled doc rows: 2=T1 3=T2 4=T3 5=T4 7=T5 8=T6 10=T7 12=T8 13=T9
#   14=T10 15=T11 16=T12 17=T13 18=T14 19=T15 (20=T16_Restricted 21=T17_MinStay
#   22=T18 23=T19) --- NOTE: APR (original T16) was dropped from stale filled

# We will work from the BLANK template (20 test rows) and rebuild properly.
# Blank template rows: 2=T1..5=T4, 7=T5, 8=T6, 10=T7, 12=T8..19=T15,
#   20=T16(APR), 21=T17(Restricted), 22=T18(MinStay), 23=T19(SpecReq), 24=T20(Taxes)

# We'll load the stale filled, fix each cell content.
# Row mapping in STALE filled:
#   row 2  → Test 1  PASS
#   row 3  → Test 2  PASS
#   row 4  → Test 3  PASS
#   row 5  → Test 4  PASS
#   row 7  → Test 5  PASS
#   row 8  → Test 6  PASS
#   row 10 → Test 7  PASS
#   row 12 → Test 8  PASS
#   row 13 → Test 9  PASS
#   row 14 → Test 10 PASS
#   row 15 → Test 11 PASS
#   row 16 → Test 12 PASS
#   row 17 → Test 13 PASS
#   row 18 → Test 14 PASS
#   row 19 → Test 15 PASS
#   row 20 → [MISSING APR TEST 16 — stale filled dropped it, replaced with T17]
#   row 20 in stale = "16. Restricted Cancellation Rules" (should be 17)
#   row 21 in stale = "17. Minimum Stay" (should be 18)
#   row 22 in stale = "18. Special Requests" (should be 19)
#   row 23 in stale = "19. Taxes & Fees" (should be 20)
#   row 24 → empty → will become Test 21 (new 2-room cancel)

# April 3 results per test number:
RESULTS = {
    1:  ("PASS",
         "Full booking flow (Flow A): searchHotels \u2192 getRooms (browse) \u2192 getRooms (blocking with status=checked validation) \u2192 confirmBooking. "
         "Tested with 2 adults, 1 night, Dubai. Booking code: 932551843. "
         "Salutation IDs dynamically mapped via getsalutationsids API. "
         "Correct adultsCode and passenger names passed."),

    2:  ("PASS",
         "Booking created for 2 adults + 1 child (age 11). Child passed with <child runno=\"0\">11</child>. "
         "All 3 passengers listed in passengersDetails. "
         "Salutation codes mapped dynamically via getsalutationsids API (no hardcoded values). "
         "Booking code: 932551903."),

    3:  ("PASS",
         "Booking created for 2 adults + 2 children (ages 8, 9). "
         "Children passed as <child runno=\"0\">8</child> and <child runno=\"1\">9</child>. "
         "All 4 passengers listed in passengersDetails. Booking confirmed successfully."),

    4:  ("PASS",
         "Multi-room booking with <rooms no=\"2\">: Room 0 = 1 adult single, Room 1 = 1 adult double. "
         "Both rooms confirmed in a single confirmBooking call. Booking code: 932552007."),

    5:  ("PASS",
         "Two-step cancellation: cancelBooking with confirm=no returns charge=0 (outside deadline, +90 days). "
         "Then cancelBooking with confirm=yes with penaltyApplied=0. "
         "Correct service code passed from step 1 response. Free cancellation verified."),

    6:  ("PASS",
         "Two-step cancellation with penalty for 2 rooms within cancellation deadline. "
         "cancelBooking confirm=no returns penalty amount. "
         "cancelBooking confirm=yes with penaltyApplied=penalty amount shown to user before confirmation. "
         "Penalty correctly displayed and confirmed."),

    7:  ("PASS",
         "After cancellation, productsLeftOnItinerary element is checked. "
         "If value > 0, application displays message that not all services have been cancelled. "
         "If value = 0, confirms complete cancellation. "
         "deleteItinerary method also implemented for APR/itinerary-based bookings."),

    8:  ("PASS",
         "getRooms request includes <roomField>tariffNotes</roomField>. "
         "tariffNotes content extracted from rateBasis and displayed in application. "
         "Content includes rate notes, hotel policies, and supplier-specific terms "
         "(e.g. \"Compulsory Tourism Dirham to be paid by guest directly at the hotel\"). "
         "Same details included on customer booking vouchers."),

    9:  ("PASS",
         "Cancellation rules sourced exclusively from getRooms response cancellationPolicy. "
         "Rules displayed with deadline dates and penalty amounts formatted for customer. "
         "Handled: free cancellation, partial penalty (percentage and fixed), and non-refundable. "
         "Displayed prominently in WhatsApp pre-booking message and web UI."),

    10: ("PASS",
         "sanitizePassengerName() enforces: minimum 2 chars, max 64 chars, letters only (no digits/special). "
         "salutation=147 used for 'Mr' (fixed from runno=1, per getsalutationsids API). "
         "salutation=14632 used for 'master'/'child' (backward compat alias added). "
         "All passenger names validated and salutation IDs fetched dynamically. Booking code: 932566283."),

    11: ("PASS",
         "Application reads totalMinimumSelling and totalMinimumSellingInRequestedCurrency from getRooms rateBasis. "
         "If MSP present, selling price is validated to be >= MSP value. "
         "MSP now propagated from totalMinimumSelling field (Phase 24 fix). "
         "MSP displayed to agents. Tested with hotel 809755 (Conrad London St James)."),

    12: ("PASS",
         "All HTTP requests include Accept-Encoding: gzip, deflate header. "
         "CURLOPT_ENCODING configured for automatic gzip/deflate decompression. "
         "Verified with getServingCountries request. "
         "Header present in every API call throughout the booking lifecycle."),

    13: ("PASS",
         "After getRooms blocking step (with roomTypeSelected), each room's rateBasis status is validated. "
         "If status != \"checked\", booking process is aborted with error: \"Rate no longer available. Please search again.\" "
         "roomField removed from blocking getRooms calls (Phase 24 fix). "
         "Multi-room validation ensures ALL rooms must have status=checked before proceeding."),

    14: ("PASS",
         "Search with 3 adults + 1 child (age 12) triggers changedOccupancy. "
         "When changedOccupancy element present: adultsCode = validForOccupancy/adults (4), "
         "actualAdults = original search (3), children no=\"0\" (child converted to adult), "
         "actualChildren no=\"1\" with actualChild runno=\"0\">12 (original child), "
         "extraBed from validForOccupancy. "
         "Variable actualAdults used throughout with verification logging. Booking code: 932552463."),

    15: ("PASS",
         "Tested with DOTW-provided hotel 2344175 (The S Hotel Al Barsha, Dubai, 14-15 May 2026, 2A+2C ages 8,12). "
         "getRooms request includes <roomField>specials</roomField>. "
         "XSD fix applied: city element removed from hotelId filter (Phase 24 fix). "
         "1 special found on roomType, specialsApplied confirmed on rateBasis. "
         "Promotions handled per rate basis and displayed to agent before booking."),

    16: ("SKIP / N/A",
         "APR flow removed from DOTW API per Olga's instruction (email 2026-03-27). "
         "Confirmed via email: \"Please disregard APRs, these have been recently removed from our API.\" "
         "Test is not applicable. Flow B (savebooking \u2192 bookitinerary) retained in codebase for legacy support only."),

    17: ("SKIP",
         "Implementation complete for cancelRestricted and amendRestricted flags. "
         "Both flags parsed from cancellationPolicy XML. "
         "When true, cancel/amend buttons disabled in UI and API calls blocked. "
         "Skipped: No hotels with cancelRestricted=true or amendRestricted=true found in DOTW sandbox inventory "
         "after scanning 163 Dubai hotels and hotel 809755 (Conrad London St James). "
         "Requesting specific hotel IDs from Olga/DOTW to complete this test."),

    18: ("PASS",
         "Implementation complete. getRooms response parsed for <minStay> and <dateApplyMinStay> elements per rateBasis. "
         "When minStay > 0, minimum night stay communicated to user. "
         "Application enforces booking nights >= minStay. "
         "Tested with DOTW sandbox hotel confirming minStay enforcement. Booking validated."),

    19: ("PASS",
         "23 correct DOTW internal special request codes wired: "
         "1711 (Non-Smoking Room), 1719 (Baby Cot/Crib), and 21 others per DOTW code list. "
         "Special request codes sent in confirmBooking XML as "
         "<specialRequests count=\"N\"><req runno=\"0\">CODE</req>...</specialRequests>. "
         "Multiple special requests supported with incremental runno. "
         "DOTW internal codes used exclusively (no free-text). Booking code: 932565683."),

    20: ("PASS",
         "getRooms response parsed for <propertyFees> array. "
         "Each fee includes name, amount, currency, and includedinprice attribute. "
         "Fees split into two categories: "
         "(1) Included in price: shown as part of total; "
         "(2) Payable at property (includedinprice=No): displayed separately with clear label. "
         "Tested: Hotel 30914, fee: \"Taxes and Fees\", includedinprice=No. "
         "All tax/fee information displayed before booking confirmation."),

    21: ("PASS",
         "NEW TEST added per CERT-06 evidence request (2-room cancellation). "
         "DotwCertify::runTest21 added: 7 steps (21a search \u2192 21b getRooms \u2192 21c blocking \u2192 "
         "21d confirm 2-room booking \u2192 21e cancelBooking confirm=no \u2192 21f cancelBooking confirm=yes \u2192 "
         "21g verify productsLeftOnItinerary=0). "
         "productsLeftOnItinerary=0 verified after both rooms cancelled. "
         "Full RQ/RS XML evidence in submission zip (dotw-certification-v2-april2026.zip). "
         "Booking codes documented in test log."),
}


def set_cell_text(cell, text):
    """Replace all text in a cell with new text, preserving first paragraph's formatting."""
    # Clear existing paragraphs except the first
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)

    # Set text in first paragraph
    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ""

    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def add_row_from_template(table, template_row_idx):
    """Copy an existing row structure and append a new row to the table."""
    # Get the table's XML element
    tbl = table._tbl
    template_row = table.rows[template_row_idx]._tr
    new_row = copy.deepcopy(template_row)
    tbl.append(new_row)
    return table.rows[-1]


# ── Load stale filled doc ────────────────────────────────────────────────────
doc = Document(SRC)
table = doc.tables[0]

print(f"Loaded stale doc. Table has {len(table.rows)} rows.")

# ── The stale filled version is MISSING the APR row (template row 20 = test 16)
# We need to insert it. The stale has:
#   row 19 = Test 15 (Special Promotions) — correct
#   row 20 = "16. Restricted Cancellation Rules" — should be test 17
#   row 21 = "17. Minimum Stay" — should be test 18
#   row 22 = "18. Special Requests" — should be test 19
#   row 23 = "19. Taxes & Fees" — should be test 20
#   row 24 = empty — will be test 21

# Step 1: Fix the test number labels and content for rows 20-23 (shift +1)
# Row 20: change "16." to "16." (APR) — we'll INSERT a new row here instead
# Actually: we'll update rows 20-23 with correct test numbers and keep stale observations,
# then use row 24 (empty) for test 21.

# Mapping: stale_row → (new_num_label, test_num_for_results)
stale_row_to_test = {
    2:  1,
    3:  2,
    4:  3,
    5:  4,
    7:  5,
    8:  6,
    10: 7,
    12: 8,
    13: 9,
    14: 10,
    15: 11,
    16: 12,
    17: 13,
    18: 14,
    19: 15,
    # Row 20 in stale = currently "16. Restricted..." — we need to insert APR before it
    # We'll handle 20-24 after insertion
}

# ── INSERT APR row (test 16) before current row 20 ──────────────────────────
# We insert a new row by copying row 20's structure and pushing it down.
# python-docx doesn't have insert_row_before, so we do XML manipulation.

tbl_elem = table._tbl
row_20 = table.rows[20]._tr  # "16. Restricted..." in stale

# Deep copy row 20 as our APR row template
apr_row_elem = copy.deepcopy(row_20)

# Insert before row 20
tbl_elem.insert(list(tbl_elem).index(row_20), apr_row_elem)

print(f"After APR row insert: Table has {len(table.rows)} rows.")

# Now the table has 27 rows:
# row 20 = APR row (newly inserted, copy of old row 20)
# row 21 = old row 20 (Restricted, label "16.") → needs relabeling to "17."
# row 22 = old row 21 (MinStay, "17.") → "18."
# row 23 = old row 22 (SpecReq, "18.") → "19."
# row 24 = old row 23 (Taxes, "19.") → "20."
# row 25 = old row 24 (empty) → test 21
# row 26 = old row 25 (footnote)

# ── Fix APR row (row 20) ────────────────────────────────────────────────────
apr_row = table.rows[20]
# Col 0: number
set_cell_text(apr_row.cells[0], "16.")
# Col 1: keep original APR test description from blank template — we'll set a summary
set_cell_text(apr_row.cells[1],
    "APRs\n"
    "Book an APR and make sure that your application does not allow cancellations or amendments. "
    "(Non-Refundable rates: nonrefundable=\"yes\")")
# Col 2: result
status_16, obs_16 = RESULTS[16]
set_cell_text(apr_row.cells[2], f"Status: {status_16}\nObservations: {obs_16}")

# ── Fix row 21 (Restricted Cancel — now correctly test 17) ─────────────────
row_21 = table.rows[21]
set_cell_text(row_21.cells[0], "17.")
status_17, obs_17 = RESULTS[17]
set_cell_text(row_21.cells[2], f"Status: {status_17}\nObservations: {obs_17}")

# ── Fix row 22 (Min Stay — now correctly test 18) ──────────────────────────
row_22 = table.rows[22]
set_cell_text(row_22.cells[0], "18.")
status_18, obs_18 = RESULTS[18]
set_cell_text(row_22.cells[2], f"Status: {status_18}\nObservations: {obs_18}")

# ── Fix row 23 (Special Requests — now correctly test 19) ──────────────────
row_23 = table.rows[23]
set_cell_text(row_23.cells[0], "19.")
status_19, obs_19 = RESULTS[19]
set_cell_text(row_23.cells[2], f"Status: {status_19}\nObservations: {obs_19}")

# ── Fix row 24 (Taxes&Fees — now correctly test 20) ────────────────────────
row_24 = table.rows[24]
set_cell_text(row_24.cells[0], "20.")
# Need to add the test description (was empty before for taxes)
set_cell_text(row_24.cells[1],
    "Taxes & Fees.\n"
    "Book any property that includes Taxes and Fees and make sure to properly display "
    "all tax and fee information, distinguishing between fees included in the price and "
    "fees payable at property.")
status_20, obs_20 = RESULTS[20]
set_cell_text(row_24.cells[2], f"Status: {status_20}\nObservations: {obs_20}")

# ── Fix row 25 (empty → test 21) ────────────────────────────────────────────
row_25 = table.rows[25]
set_cell_text(row_25.cells[0], "21.")
set_cell_text(row_25.cells[1],
    "2-Room Cancellation (Additional Evidence)\n"
    "Create a booking for 2 rooms, then cancel both rooms and verify that "
    "productsLeftOnItinerary=0 is returned after full cancellation. "
    "This test provides explicit evidence of the cancelbooking process for multi-room bookings "
    "as requested per CERT-06 evidence requirement.")
status_21, obs_21 = RESULTS[21]
set_cell_text(row_25.cells[2], f"Status: {status_21}\nObservations: {obs_21}")

# ── Update existing test rows (1-15, already correct) ───────────────────────
# Only update the Status/Observations column (col 2) for tests 1-15
# since some had minor changes (test 10 salutation fix, test 11 MSP, test 13 roomField, test 15 XSD)

update_rows = {
    2:  1,
    3:  2,
    4:  3,
    5:  4,
    7:  5,
    8:  6,
    10: 7,
    12: 8,
    13: 9,
    14: 10,
    15: 11,
    16: 12,
    17: 13,
    18: 14,
    19: 15,
}

for row_idx, test_num in update_rows.items():
    row = table.rows[row_idx]
    status, obs = RESULTS[test_num]
    set_cell_text(row.cells[2], f"Status: {status}\nObservations: {obs}")

# ── Update paragraph text in body for summary counts ────────────────────────
# Check paragraphs for any summary mentions
for i, para in enumerate(doc.paragraphs):
    txt = para.text
    if "17 PASS" in txt or "17/19" in txt:
        for run in para.runs:
            run.text = run.text.replace("17 PASS", "19 PASS").replace("17/19", "19/21")
        print(f"Updated paragraph {i}: summary count")
    if "2026-03-24" in txt or "March 24" in txt or "March 2026" in txt:
        for run in para.runs:
            run.text = (run.text
                .replace("2026-03-24", "2026-04-03")
                .replace("March 24, 2026", "April 3, 2026")
                .replace("March 2026", "April 2026"))
        print(f"Updated paragraph {i}: date")

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(DEST)
print(f"\nSaved: {DEST}")
print(f"Final table rows: {len(table.rows)}")

# Print final summary
print("\nFinal table structure:")
for ri, row in enumerate(table.rows):
    no = row.cells[0].text.strip()[:20]
    status_col = row.cells[2].text.strip()[:60]
    print(f"  Row {ri}: [{no}] {status_col}")
