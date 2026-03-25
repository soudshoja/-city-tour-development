"""
Fill the DOTW XML CERTIFICATION TEST PLAN docx with latest results.
DOTW's template has 19 tests. Our internal suite has 20.
Mapping: our 16+17 -> DOTW 16, our 18->17, our 19->18, our 20->19.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

src = "docs/XML CERTIFICATION TEST PLAN_.docx"
dst = "docs/dotw-certification-submission/XML CERTIFICATION TEST PLAN_FILLED.docx"

# Latest results mapped to DOTW's 19-test numbering
results = {
    1: {
        "status": "PASS",
        "obs": "Full booking flow (Flow A): searchHotels \u2192 getRooms (browse) \u2192 getRooms (blocking with status=checked validation) \u2192 confirmBooking. Tested with 2 adults, 1 night, Dubai. Booking code: 932551843. Salutation IDs dynamically mapped via getsalutationsids API. Correct adultsCode and passenger names passed."
    },
    2: {
        "status": "PASS",
        "obs": "Booking created for 2 adults + 1 child (age 11). Child passed with <child runno=\"0\">11</child>. All 3 passengers listed in passengersDetails. Salutation codes mapped dynamically via getsalutationsids API (no hardcoded values). Booking code: 932551903."
    },
    3: {
        "status": "PASS",
        "obs": "Booking created for 2 adults + 2 children (ages 8, 9). Children use separate runno attributes: <child runno=\"0\">8</child> and <child runno=\"1\">9</child>. All 4 passengers in passengersDetails with dynamic salutation IDs. Booking code: 932551953."
    },
    4: {
        "status": "PASS",
        "obs": "Multi-room booking with <rooms no=\"2\">: Room 0 = 1 adult (single), Room 1 = 2 adults (double). Both rooms validated with status=checked before confirmBooking. Each room has separate roomTypeSelected with individual allocationDetails. Booking code: 932552043."
    },
    5: {
        "status": "PASS",
        "obs": "Two-step cancellation: cancelBooking with confirm=no returns charge=0 (outside deadline, +90 days). Then cancelBooking with confirm=yes with penaltyApplied=0. Correct service code passed from step 1 response. Free cancellation verified."
    },
    6: {
        "status": "PASS",
        "obs": "Two-step cancellation with penalty for 2 rooms within deadline. Code correctly checks charge > 0 and passes penaltyApplied value from <charge> element (NOT from <formatted> tag as per DOTW requirement). Both rooms cancelled with correct penalty amounts."
    },
    7: {
        "status": "PASS",
        "obs": "After cancellation, productsLeftOnItinerary element is checked. If value > 0, application displays message that not all services have been cancelled. If value = 0, confirms complete cancellation. deleteItinerary method also implemented for APR/itinerary-based bookings."
    },
    8: {
        "status": "PASS",
        "obs": "getRooms request includes <roomField>tariffNotes</roomField>. tariffNotes content extracted from rateBasis and displayed in application. Content includes rate notes, hotel policies, and supplier-specific terms (e.g. \"Compulsory Tourism Dirham to be paid by guest directly at the hotel\"). Same details included on customer booking vouchers."
    },
    9: {
        "status": "PASS",
        "obs": "Cancellation rules sourced exclusively from getRooms response (not searchHotels). getRooms request includes <roomField>cancellation</roomField>. Rules parsed from <cancellationRules><rule> array with fromDate, toDate, and cancelCharge per rule. Tiered penalties supported. No additional buffer applied to DOTW cancellation deadlines."
    },
    10: {
        "status": "PASS",
        "obs": "sanitizePassengerName() enforces: minimum 2 chars, maximum 25 chars, no whitespace (multi-word names merged: \"James Lee\" \u2192 \"JamesLee\"), no numbers or special characters (\"O'Brien\" \u2192 \"OBrien\"). Duplicate passenger names rejected. All passengers including children passed in confirmBooking/saveBooking."
    },
    11: {
        "status": "PASS",
        "obs": "Application reads totalMinimumSelling and totalMinimumSellingInRequestedCurrency from getRooms rateBasis. If MSP present, selling price is validated to be >= MSP value. MSP displayed to agents. Note: Our platform operates primarily as B2B for travel agencies, not direct B2C. Tested with hotel 809755 (Conrad London St James)."
    },
    12: {
        "status": "PASS",
        "obs": "All HTTP requests include Accept-Encoding: gzip, deflate header. CURLOPT_ENCODING configured for automatic gzip decompression. Verified with getServingCountries request. Header present in every API call throughout the booking lifecycle."
    },
    13: {
        "status": "PASS",
        "obs": "After getRooms blocking step (with roomTypeSelected), each room's rateBasis status is validated. If status != \"checked\", booking process is aborted with error: \"Rate no longer available. Please search again.\" Multi-room validation ensures ALL rooms must have status=checked before proceeding."
    },
    14: {
        "status": "PASS",
        "obs": "Search with 3 adults + 1 child (age 12) triggers changedOccupancy. When changedOccupancy element present: adultsCode = validForOccupancy/adults (4), actualAdults = original search (3), children no=\"0\" (child converted to adult), actualChildren no=\"1\" with actualChild runno=\"0\">12 (original child), extraBed from validForOccupancy. Booking code: 932552463."
    },
    15: {
        "status": "PASS",
        "obs": "Tested with DOTW-provided hotel 2344175 (The S Hotel Al Barsha, Dubai, 14-15 May 2026, 2A+2C ages 8,12). getRooms request includes <roomField>specials</roomField>. 1 special found on roomType, specialsApplied confirmed on rateBasis. Promotions handled per rate basis (not per room type) and displayed to agent before booking."
    },
    16: {
        "status": "SKIP - Awaiting hotel IDs from DOTW",
        "obs": "Implementation complete for both APR and Restricted Cancellation scenarios.\n\nAPR / Non-Refundable Rates: Rates with nonrefundable=\"yes\" detected. APR bookings routed to Flow B: saveBooking \u2192 bookItinerary. Cancel/amend UI disabled for APR bookings.\n\ncancelRestricted / amendRestricted: Both flags parsed from cancellation rules. When true, cancel/amend buttons disabled in UI.\n\nSkipped: No nonrefundable rates or cancelRestricted/amendRestricted flags found after scanning 163 Dubai hotels and hotel 809755 (Conrad London St James). Requesting specific hotel IDs from DOTW."
    },
    17: {
        "status": "SKIP - Awaiting hotel ID from DOTW",
        "obs": "Implementation complete. getRooms response parsed for <minStay> and <dateApplyMinStay> elements per rateBasis. When minStay > 0, minimum night stay communicated to user. Application enforces booking nights >= minStay. Skipped: No minStay values found after scanning 305 Dubai hotels and hotel 809755. Requesting specific hotel ID from DOTW, or DOTW can filter these rates server-side."
    },
    18: {
        "status": "PASS",
        "obs": "Special request code 1 (No Smoking) sent in confirmBooking XML as <specialRequests count=\"1\"><req runno=\"0\">1</req></specialRequests>. Multiple special requests supported with incremental runno. DOTW internal codes used. Booking code: 932565683."
    },
    19: {
        "status": "PASS",
        "obs": "getRooms response parsed for <propertyFees> array. Each fee includes name, amount, currency, and includedinprice attribute. Fees with includedinprice=\"No\" displayed as \"Payable at property\". Tested: Hotel 30914, fee: \"Taxes and Fees\", includedinprice=No. All tax/fee information displayed before booking confirmation."
    },
}

doc = Document(src)

# Find the table in the document
tables = doc.tables
print(f"Found {len(tables)} table(s) in document")

if tables:
    table = tables[0]
    print(f"Table has {len(table.rows)} rows and {len(table.columns)} columns")

    # The table structure is: No | Test | Status/Observations
    # Row 0 = header, then test rows with merged cells for section headers

    # Map row content to test numbers
    current_test = 0
    for i, row in enumerate(table.rows):
        cells = row.cells
        cell_texts = [c.text.strip() for c in cells]

        # Check if first cell has a test number
        first = cell_texts[0] if cell_texts else ""

        # Try to extract test number
        test_num = None
        if first and first.replace(".", "").isdigit():
            test_num = int(first.replace(".", ""))

        if test_num and test_num in results:
            # Find the Status/Observations cell (last cell)
            status_cell = cells[-1]

            # Clear existing content
            for p in status_cell.paragraphs:
                for run in p.runs:
                    run.text = ""

            # Set new content
            r = results[test_num]
            status_text = f"Status: {r['status']}"
            obs_text = f"Observations: {r['obs']}"

            # Write to first paragraph
            if status_cell.paragraphs:
                p = status_cell.paragraphs[0]
                p.clear()

                # Add status in bold
                run_status = p.add_run(status_text)
                run_status.bold = True
                run_status.font.size = Pt(9)
                if "PASS" in r["status"]:
                    run_status.font.color.rgb = RGBColor(0, 128, 0)
                elif "SKIP" in r["status"]:
                    run_status.font.color.rgb = RGBColor(255, 140, 0)

                # Add newline + observations
                run_obs = p.add_run(f"\n{obs_text}")
                run_obs.font.size = Pt(9)

            print(f"  Test {test_num}: {r['status']}")

doc.save(dst)
print(f"\nSaved to: {dst}")
print("Done!")
